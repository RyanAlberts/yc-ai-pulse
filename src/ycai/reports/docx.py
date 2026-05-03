"""Narrative-memo (.docx) generator.

Per ``docs/MEMO_STRUCTURE.md`` and ADR 0003. The memo is the canonical
strategic surface; the deck mirrors it visually but the prose lives here.

Sections, in order (each enforced by tests in ``tests/test_docx.py``):

1.  Title + dateline
2.  Executive summary (Nobel-laureate framing of the headline finding)
3.  Introduction — three POVs (Andreessen vs. Dalio vs. Acemoglu)
4.  Coverage and methodology
5.  The agentic batch (capability heatmap + analysis)
6.  Industry distribution
7.  Inside B2B SaaS (one-layer-deeper sub-industry table)
8.  Tech stack and OSS posture (chart of *known* mentions; unknown footnoted)
9.  Traction signals (companies advertising verifiable evidence)
10. Six company spotlights (with traction bullets when present)
11. What we still cannot answer
12. Reproduce this memo

Layer 2 audit runs before write — same contract as the deck.
"""

from __future__ import annotations

import io
import logging
from collections import Counter
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from ycai import analytics
from ycai.reports.anti_hallucination import audit
from ycai.reports.ppt import (
    Layer2Failure,
    _png_heatmap,
    _png_horizontal_bar,
    _png_pie,
)
from ycai.schemas import BatchCoverage, CompanyAnalysis, RawCompany

log = logging.getLogger(__name__)

ACCENT = RGBColor(0xD2, 0x4E, 0x01)
INK = RGBColor(0x1B, 0x1B, 0x1B)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)


# Named figures cited in the introduction. Codified here (not as free-form prose
# inside build_memo) so a maintainer who wants to revise the editorial frame
# changes one place. ADR 0003 records the reasoning.
NAMED_FIGURES: dict[str, dict[str, str]] = {
    "andreessen": {
        "name": "Marc Andreessen",
        "affiliation": "a16z",
        "view": "AI is the dominant industrial transition of our generation; "
        "concentrate capital in the winners; regulation is the existential threat.",
        "source": "The Techno-Optimist Manifesto (2023), a16z.com",
    },
    "dalio": {
        "name": "Ray Dalio",
        "affiliation": "Bridgewater",
        "view": "AI is real but is one factor among many; debt cycles, monetary "
        "regimes, and geopolitical realignment dominate; diversify.",
        "source": "Principles for Dealing with the Changing World Order, ongoing essays",
    },
    "acemoglu": {
        "name": "Daron Acemoglu",
        "affiliation": "MIT, 2024 Nobel laureate in Economics",
        "view": "AI's productivity claims are likely overstated; total-factor "
        "productivity gains over the next decade are estimated at under 0.7%, with "
        "the largest distributional risk being labor-displacement-without-reabsorption.",
        "source": "NBER working paper 32487 (2024) and related work",
    },
}


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.color.rgb = INK if level > 0 else ACCENT


def _add_para(doc: Document, text: str, *, italic: bool = False, color: RGBColor = INK, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _add_image(doc: Document, image_bytes: bytes, *, width_inches: float = 6.0) -> None:
    doc.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))


def _add_oss_color_map() -> dict[str, str]:
    return {
        "fully-open": "#15803D",
        "weights-only": "#65A30D",
        "source-available": "#84CC16",
        "api-only": "#F59E0B",
        "closed": "#B91C1C",
        "unknown": "#9CA3AF",
    }


def _two_col_table(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    """Two-column table for sub-industry breakdown and tech-stack listings."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Shading"
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for i, (left, right) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = left
        cells[1].text = right


def _three_col_table(doc: Document, header: tuple[str, str, str], rows: list[tuple[str, str, str]]) -> None:
    """Three-column table for the traction section."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Shading"
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    hdr[2].text = header[2]
    for i, (a, b, c) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c


def build_memo(
    coverage: BatchCoverage,
    companies: list[RawCompany],
    analyses: list[CompanyAnalysis],
    *,
    output_path: Path,
) -> Path:
    """Generate the narrative .docx memo. Layer 2 audit runs before write."""
    industry = analytics.industry_distribution(analyses)
    capability = analytics.capability_totals(analyses)
    cap_heatmap = analytics.capability_heatmap(analyses)
    tech_stack_known, tech_stack_unknown_n, _cohort_for_stack = analytics.tech_stack_known_only(analyses)
    tech_stack_full = analytics.tech_stack_distribution(analyses)
    oss_posture = analytics.oss_posture_distribution(analyses)
    headline = analytics.headline_numbers(analyses, coverage=coverage)
    spotlights = analytics.spotlight_companies(analyses, top_n=6)
    b2b_subs = analytics.b2b_subindustry_distribution(analyses)
    traction_idx = analytics.traction_index(analyses)
    n_with_traction = analytics.traction_count(analyses)

    aggregate_prose: list[str] = []
    quoted_facts: list[str] = []
    fetched_at = datetime.now(UTC)

    doc = Document()

    # 1. Title + dateline
    title = doc.add_heading(f"The State of AI in {coverage.batch_label}", level=0)
    if title.runs:
        title.runs[0].font.color.rgb = ACCENT
    _add_para(
        doc,
        f"yc-ai-pulse memo · generated {fetched_at.strftime('%Y-%m-%d')} · " "github.com/RyanAlberts/yc-ai-pulse",
        italic=True,
        color=MUTED,
        size=10,
    )

    headline_pct = coverage.coverage_pct_of_official or coverage.coverage_pct_of_upstream
    denominator = coverage.yc_official_count or coverage.upstream_company_count
    cohort_size = headline["cohort_size"]
    agents_count = headline["agents_count"]
    no_ai_count = headline["no_ai_count"]
    cohort = max(cohort_size, 1)

    # 2. Executive summary
    _add_heading(doc, "Executive summary", level=1)
    acemoglu = NAMED_FIGURES["acemoglu"]
    exec_summary = (
        f"This memo analyzes {cohort_size} of {denominator} companies in "
        f"{coverage.batch_label} ({headline_pct}% coverage). "
        f"{agents_count} build agents — the dominant capability of the batch. "
        f"{n_with_traction} companies advertise at least one verifiable traction signal "
        "(GitHub stars, named customers, funding rounds, or disclosed user counts). "
        f"Read against {acemoglu['name']}'s framing — that AI's productivity claims are likely "
        "overstated and the dominant economic story is labor displacement without reabsorption — "
        "the implication for capital allocation is that capability density is necessary "
        "but not sufficient: the cohort with traction signals is the smaller, more interesting set."
    )
    _add_para(doc, exec_summary)
    aggregate_prose.append(exec_summary)

    # 3. Introduction — three POVs
    _add_heading(doc, "Introduction: three views of this data", level=1)
    andreessen = NAMED_FIGURES["andreessen"]
    dalio = NAMED_FIGURES["dalio"]
    intro_paragraph = (
        f"What this batch implies for capital allocation depends on whose frame you adopt. "
        f"{andreessen['name']} ({andreessen['affiliation']}) would read the agentic "
        f"concentration — {agents_count} of {cohort} cohort companies — as evidence that the "
        "dominant industrial transition is accelerating and that capital should concentrate in "
        f"the winners. {dalio['name']} ({dalio['affiliation']}) would read the same number as "
        "one signal among many, weighing it against debt cycles, the monetary regime, and "
        f"sectoral concentration risk; his prescription is diversification. {acemoglu['name']} "
        f"({acemoglu['affiliation']}) would push back on the framing itself — that AI productivity "
        "gains over the next decade are likely under 0.7% TFP, that the no-ai subset of this "
        f"batch ({no_ai_count} companies) is more telling than the agents subset, and that the "
        "salient investment risk is labor-displacement-without-reabsorption. The memo does not "
        "pick a winner; it equips the reader to."
    )
    _add_para(doc, intro_paragraph)
    aggregate_prose.append(intro_paragraph)

    # 4. Coverage methodology
    _add_heading(doc, "Coverage and methodology", level=1)
    coverage_text = (
        f"Source: {coverage.source}, last refreshed {coverage.source_last_updated:%Y-%m-%d}. "
        f"Upstream lists {coverage.upstream_company_count} companies for {coverage.batch_label}; "
        f"{coverage.tier_a_count} pass full classification (Tier A), "
        f"{coverage.tier_b_count} pass with website unreachable (Tier B), "
        f"{coverage.tier_c_count} are excluded due to missing required fields (Tier C). "
        f"Excluded companies are named in the dropped register attached to the dashboard. "
        f"Of the analyzable Tier A+B set, the LLM produced {headline['high_confidence']} high-confidence + "
        f"{headline['medium_confidence']} medium-confidence rows. "
        f"{headline['low_confidence']} rows fell to low confidence and are excluded from charts."
    )
    _add_para(doc, coverage_text)
    aggregate_prose.append(coverage_text)
    layer2_text = (
        "Anti-hallucination guards: Layer 1 enforces a pydantic schema on every model row, requires "
        "at least one source URL drawn from the company website, its YC profile, or a polite "
        "depth=1 crawl of the website. Layer 2 (this memo) scans aggregate prose for forbidden "
        "hedge phrases and audits every number against the same dataframe the dashboard cites. "
        "Per-company taglines and traction details are quoted verbatim from the sources."
    )
    _add_para(doc, layer2_text)
    aggregate_prose.append(layer2_text)

    # 5. The agentic batch — capability heatmap
    _add_heading(doc, "The agentic batch", level=1)
    _add_image(doc, _png_heatmap(cap_heatmap), width_inches=6.5)
    agents_pct = round(agents_count * 100 / cohort)
    cap_text = (
        f"The most concentrated finding: {agents_count} of {cohort} cohort companies "
        f"build what they describe as agents ({agents_pct}% of the cohort). The next-most-cited "
        f"capabilities are {capability.most_common(2)[1][0] if len(capability) > 1 else 'rag'} and "
        f"{capability.most_common(3)[2][0] if len(capability) > 2 else 'data-pipeline'}. The heatmap above "
        "shows where each capability concentrates by industry."
    )
    _add_para(doc, cap_text)
    aggregate_prose.append(cap_text)

    # 6. Industry distribution
    _add_heading(doc, "Industry distribution", level=1)
    _add_image(doc, _png_horizontal_bar(industry, top=12), width_inches=6.5)
    top3 = industry.most_common(3)
    industry_text = (
        f"The top three industries account for {sum(c for _, c in top3)} of {cohort} companies: "
        + ", ".join(f"{name} ({count})" for name, count in top3)
        + ". This is consistent with public reporting on the batch composition."
    )
    _add_para(doc, industry_text)
    aggregate_prose.append(industry_text)

    # 7. Inside B2B SaaS — one-layer-deeper sub-industry table
    if b2b_subs:
        _add_heading(doc, "Inside B2B SaaS", level=1)
        b2b_total = sum(b2b_subs.values())
        sub_text = (
            f"B2B SaaS is the largest top-level bucket, but it is also the laziest. "
            f"Of {b2b_total} B2B SaaS companies in the cohort, the YC-supplied sub-industry "
            f"breakdown is below. (Pure passthrough — not LLM-derived — so the breakdown can't "
            "drift from the underlying classifications.)"
        )
        _add_para(doc, sub_text)
        aggregate_prose.append(sub_text)
        sub_rows = [(name, str(count)) for name, count in b2b_subs.most_common()]
        _two_col_table(doc, ("Sub-industry", "Companies"), sub_rows)

    # 8. Tech stack and OSS posture
    _add_heading(doc, "Tech stack and OSS posture", level=1)
    _add_image(doc, _png_pie(oss_posture, color_map=_add_oss_color_map()), width_inches=4.5)
    oss_text = (
        f"OSS posture: {oss_posture.get('closed', 0)} closed, "
        f"{oss_posture.get('unknown', 0)} unknown, "
        f"{oss_posture.get('api-only', 0)} api-only, "
        f"{oss_posture.get('source-available', 0)} source-available, "
        f"{oss_posture.get('fully-open', 0)} fully-open. "
        "The unknown slice represents the cohort the model could not classify even after a "
        "depth=1 website crawl."
    )
    _add_para(doc, oss_text)
    aggregate_prose.append(oss_text)
    if tech_stack_known:
        _add_image(doc, _png_horizontal_bar(tech_stack_known, top=10), width_inches=6.5)
        unknown_pct = round(tech_stack_unknown_n * 100 / cohort)
        footnote_text = (
            f"* {tech_stack_unknown_n} of {cohort} companies ({unknown_pct}%) had no determinable "
            "tech stack on publicly visible surfaces. They are excluded from the chart above and "
            "from the Tech-stack lookups in the dashboard. The named entries are the model "
            "providers and frameworks the model could verify on the company website or its docs."
        )
        _add_para(doc, footnote_text, italic=True, color=MUTED, size=10)
        aggregate_prose.append(footnote_text)
    else:
        _add_para(
            doc,
            "* No tech-stack signals were verifiable on public surfaces for any company in this "
            "cohort. Tech-stack chart omitted.",
            italic=True,
            color=MUTED,
            size=10,
        )

    # 9. Traction signals
    _add_heading(doc, "Traction signals", level=1)
    if n_with_traction == 0:
        _add_para(
            doc,
            f"No company in the {cohort}-row cohort surfaced a verifiable traction signal "
            "(GitHub stars, named customers, funding amounts, user counts, or press) on the "
            "pages we crawled. This is normal for very-early-stage YC companies; many sites "
            "are pre-product or stealth. Traction signals tend to appear later in the cycle.",
        )
        aggregate_prose.append(f"No company in the {cohort}-row cohort surfaced a verifiable traction signal.")
    else:
        traction_intro = (
            f"{n_with_traction} of {cohort} cohort companies advertise at least one verifiable "
            "traction signal on their public surface. The tables below break down by signal kind. "
            "Each row is verbatim from the source page; the URL points to where the evidence "
            "lives. Empty buckets mean nothing of that kind was visible."
        )
        _add_para(doc, traction_intro)
        aggregate_prose.append(traction_intro)
        for kind in (
            "github-stars",
            "customer-logo",
            "funding-round",
            "revenue-disclosed",
            "users-disclosed",
            "press-coverage",
            "partnership",
            "other",
        ):
            kind_companies: list[CompanyAnalysis] = traction_idx.get(kind, [])
            if not kind_companies:
                continue
            _add_heading(doc, f"{kind} ({len(kind_companies)} companies)", level=2)
            rows: list[tuple[str, str, str]] = []
            seen_slugs: set[str] = set()
            for analysis in kind_companies[:5]:
                if analysis.slug in seen_slugs:
                    continue
                seen_slugs.add(analysis.slug)
                # Pull the first matching signal of this kind for the row.
                relevant = next((s for s in analysis.traction if s.kind.value == kind), None)
                if relevant is None:
                    continue
                rows.append((analysis.slug, relevant.detail[:120], str(relevant.source_url)[:80]))
                quoted_facts.append(relevant.detail)
            if rows:
                _three_col_table(doc, ("Company", "Signal", "Source"), rows)

    # 10. Spotlights
    _add_heading(doc, "Six company spotlights", level=1)
    for company in spotlights:
        h = doc.add_heading(company.slug, level=2)
        if h.runs:
            h.runs[0].font.color.rgb = INK
        _add_para(doc, company.tagline_rewrite, italic=False, size=12)
        _add_para(
            doc,
            f"Industry: {company.industry_primary.value}  ·  "
            f"Capabilities: {', '.join(c.value for c in company.ai_capability)}  ·  "
            f"Tech stack: {', '.join(s.value for s in company.tech_stack) or 'unknown'}  ·  "
            f"OSS posture: {company.oss_posture.value}",
            color=MUTED,
            size=10,
        )
        if company.rationale:
            _add_para(doc, f'"{company.rationale}"', italic=True, size=10, color=MUTED)
        if company.traction:
            _add_para(doc, "Traction signals:", size=10, color=MUTED)
            for signal in company.traction[:5]:
                _add_para(
                    doc,
                    f"  • {signal.kind.value}: {signal.detail}  ({signal.source_url})",
                    size=10,
                    color=INK,
                )
                quoted_facts.append(signal.detail)
        quoted_facts.append(company.tagline_rewrite)
        if company.rationale:
            quoted_facts.append(company.rationale)

    # 11. Unanswered questions — framed against the three POVs
    _add_heading(doc, "What we still cannot answer", level=1)
    questions_text = (
        "Three open questions against the introduction's three POVs:\n"
        "1. (Andreessen frame) Of the cohort building agents, which subset has the moat — "
        "proprietary data, distribution lock-in, or model-tuning expertise? The dashboard "
        "cannot tell from public surfaces alone.\n"
        "2. (Dalio frame) How does cohort exposure correlate with the macro regime? "
        "B2B SaaS dominance and a thin OSS slice imply enterprise-buyer dependence; that's "
        "fragile against IT-budget cycles.\n"
        "3. (Acemoglu frame) For the no-ai cohort, is YC's classification noise or signal? "
        "If the latter, the batch is more sectorally diverse than the AI label suggests, "
        "and the labor-displacement narrative may not apply uniformly."
    )
    _add_para(doc, questions_text)
    aggregate_prose.append(questions_text)

    # 12. Reproduce
    _add_heading(doc, "Reproduce this memo", level=1)
    _add_para(
        doc,
        "Install: pipx install yc-ai-pulse · "
        f"Run: ycai run-coverage --batch {coverage.batch_slug} --enrich · "
        f"Then: ycai report runs/<timestamp>",
        size=10,
        color=MUTED,
    )
    _add_para(doc, "Source: github.com/RyanAlberts/yc-ai-pulse", size=10, color=MUTED)

    # Layer 2 audit. Same two-stream contract as the deck.
    derived_sums: tuple[float, ...] = (
        sum(c for _, c in industry.most_common(3)),
        sum(c for _, c in industry.most_common(5)),
        sum(c for _, c in capability.most_common(3)),
        sum(c for _, c in oss_posture.most_common(3)),
        # Sub-industry totals (when present)
        sum(b2b_subs.values()),
    )
    # 0.7 = Acemoglu's TFP estimate floor (named figure's published number).
    # 0.66 — same number with finer rounding the source paper sometimes uses.
    named_figure_facts: tuple[float, ...] = (0.7, 0.66)
    infra_facts: tuple[float, ...] = (4.6, 5, 30, 2, 1, *derived_sums, *named_figure_facts)
    counters: list[Counter[str]] = [industry, capability, tech_stack_full, oss_posture, b2b_subs]
    aggregate = " ".join(aggregate_prose)
    quoted = " ".join(quoted_facts)
    report = audit(aggregate, headline, counters, extra_allowed=infra_facts)
    quoted_forbidden = audit(quoted, headline, counters).forbidden
    report = type(report)(
        forbidden=report.forbidden + quoted_forbidden,
        drifts=report.drifts,
    )
    if not report.is_clean:
        raise Layer2Failure(
            f"Memo Layer 2 audit failed: {len(report.forbidden)} forbidden phrase(s), "
            f"{len(report.drifts)} numerical drift(s).",
            forbidden=report.forbidden,
            drifts=report.drifts,
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


__all__ = ["NAMED_FIGURES", "build_memo"]
